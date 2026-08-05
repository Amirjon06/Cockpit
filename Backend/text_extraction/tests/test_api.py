from io import BytesIO

from docx import Document
from fastapi.testclient import TestClient

from api import app


def test_extract_endpoint_returns_text_for_docx():
    document = Document()
    document.add_paragraph("Routing Basics")
    document.add_paragraph("Routers choose paths using routing tables.")
    buffer = BytesIO()
    document.save(buffer)

    client = TestClient(app)
    response = client.post(
        "/extract",
        files={
            "file": (
                "routing.docx",
                buffer.getvalue(),
                (
                    "application/vnd.openxmlformats-officedocument."
                    "wordprocessingml.document"
                ),
            )
        },
    )

    assert response.status_code == 200
    assert response.json()["text"] == (
        "Routing Basics\n\nRouters choose paths using routing tables."
    )


def test_extract_endpoint_returns_clear_error_for_unsupported_file():
    client = TestClient(app)
    response = client.post(
        "/extract",
        files={"file": ("archive.zip", b"not a document", "application/zip")},
    )

    assert response.status_code == 400
    assert "Unsupported file type" in response.json()["detail"]
